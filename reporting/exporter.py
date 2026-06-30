import os 
import pandas as pd
from docx import Document
from docx.shared import Pt

from .utils import (format_descriptive_table, format_posthoc_table, format_p)
from .report import build_anova_summary_table


def export_anova_to_csv(anova_results, output_dir):
    """
    Export ANOVA results into CSV files (one file per DV).
    """

    os.makedirs(output_dir, exist_ok=True)

    for item in anova_results:
        dv = item["dv"]
        sections = []

        # Table 1: Wave Mean Trends
        wave_means = item["wave_means"].copy().reset_index()
        wave_means.insert(0, "Section", "Table 1 - Wave Mean Trends")
        sections.append(wave_means)

        # Table 2-4: Per Wave
        for wave, result in item["by_wave"].items():
        
            # Group stats
            group_stats = format_descriptive_table(result.group_stats).copy()
            group_stats.rename(columns={"Std. Error": "Std. Error (Group Stats)"}, inplace=True)
            group_stats.insert(0, "Section", f"{wave} - Group Statistics")
            sections.append(group_stats)

            # ANOVA summary
            anova_summary = build_anova_summary_table(result).copy()
            anova_summary.insert(0, "Section", f"{wave} - ANOVA Summary")
            sections.append(anova_summary)

            # Post-hoc (if exists)
            if result.is_significant and result.posthoc_df is not None:
                posthoc = format_posthoc_table(result.posthoc_df).copy()

                posthoc = posthoc.drop(
                    columns=[
                        "Tukey_Statistic",
                        "Effect size (Hedges’ g)"
                    ],
                    errors="ignore"
                )

                posthoc.rename(columns={
                    "Std. Error": "Std. Error (Post-hoc)"
                }, inplace=True)

                posthoc.insert(
                        0,
                        "Section",
                        f"{wave} - Post-hoc"
                    )

                sections.append(posthoc)

        final_df = pd.concat(
            sections,
            ignore_index=True,
            sort=False
        )

        file_path = os.path.join(output_dir, f"{dv}.csv")
        final_df.to_csv(file_path, index=False)

        print(f"Saved CSV: {file_path}")


def df_to_word_table(doc, df):
    """
    Convert DataFrame to Word table
    """
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"

    # header
    for i, col in enumerate(df.columns):
        table.rows[0].cells[i].text = str(col)

    # rows
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = "" if pd.isna(val) else str(val)


def add_spacing(doc, pt=12):
    """
    Add vertical space in Word document.
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pt)


def export_anova_to_docx(anova_results, output_dir):
    """
    Export ANOVA results into a Word report (one file per DV).
    """
    os.makedirs(output_dir, exist_ok=True)

    for item in anova_results:
        dv = item["dv"]
        dv_title = dv.replace("_", " ").title()

        doc = Document()

        doc.add_heading(f"One-Way ANOVA Report: {dv_title}", level=1)
        add_spacing(doc, 12)

        doc.add_heading("Overview Across Waves", level=1)

        doc.add_heading(f"Table 1. {dv_title} Index Mean Trends by Wave", level=2)
        wave_means = item["wave_means"].copy().reset_index()
        wave_means = wave_means.round(2)
        df_to_word_table(doc, wave_means)
        add_spacing(doc, 18)

        doc.add_heading(f"Detailed ANOVA Results by Wave", level=1)

        for wave, result in item["by_wave"].items():
            doc.add_heading(f"Wave: {wave}", level=2)
            doc.add_heading(f"ANOVA Summary", level=2)

            eta_str = (
                f", Eta Squared = {result.eta_squared:.3f}"
                if result.eta_squared is not None
                else ""
            )
            doc.add_paragraph(
                f"A one-way ANOVA was conducted to examine differences in {dv_title} across groups for {wave}."
            )
            doc.add_paragraph(
                f"Results indicated a "
                f"{'statistically significant' if result.is_significant else 'non-significant'} "
                f"effect of group on {dv_title}, "
                f"F({result.df_between}, {result.df_within}) = {result.F:.2f}, "
                f"{format_p(result.p)}{eta_str}."
            )

            add_spacing(doc, 12)

            # Group Statistics
            doc.add_heading(f"Table 2. Clusters and {dv_title} ANOVA", level=2)

            doc.add_paragraph(
                f"{dv_title} Index Mean Values\n"
                f"F = {result.F:.2f} (difference between groups), {format_p(result.p)}"
            )

 
            group_stats = format_descriptive_table(result.group_stats).copy()
            group_stats.rename(
                columns={"Std. Error": "Std. Error (Group Stats)"},
                inplace=True
            )
            df_to_word_table(doc, group_stats)

            add_spacing(doc, 12)

            # ANOVA Summary
            doc.add_heading(f"Table 3. One-Way ANOVA Results", level=2)

            anova_summary = build_anova_summary_table(result).copy()
            df_to_word_table(doc, anova_summary)

            add_spacing(doc, 12)

            # POSTHOC
            if result.is_significant and result.posthoc_df is not None:

                doc.add_heading(f"Table 4. Post-hoc Comparisons (Tukey HSD)", level=2)

                posthoc = format_posthoc_table(result.posthoc_df).copy()

                posthoc = posthoc.drop(
                    columns=[
                        "Tukey_Statistic",
                        "Effect size (Hedges’ g)"
                    ],
                    errors="ignore"
                )

                posthoc.rename(
                    columns={"Std. Error": "Std. Error (Post-hoc)"},
                    inplace=True
                )

                df_to_word_table(doc, posthoc)

            add_spacing(doc, 18)


        file_path = os.path.join(output_dir, f"{dv}.docx")
        doc.save(file_path)

        print(f"Saved DOCX: {file_path}")


